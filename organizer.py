"""
Oracle Document Organizer
=========================
Scans an _Inbox/ folder, classifies each file by Oracle module and version,
renames it using a consistent convention, and moves it to the right folder.

Folder structure created automatically:
  <root>/
    _Inbox/                   <- drop new files here
    EPM/
      FCCS/ ARCS/ PBCS/ Free_Form/ Other/
    R12_ERP/
      AP/ AR/ GL/ CE/ AHCS_FAH/ PO/ Security/
      CST/ PA/ FA/ Other/
    Fusion_Cloud_ERP/
      AP/ AR/ GL/ CE/ AHCS_FAH/ PO/ Security/
      CST/ PA/ FA/ Other/
    Other/
      Data_Models/            <- SQL / data model files
      Other/                  <- unclassified

File naming convention:
  MODULE_VERSION_Title_YYYY-MM-DD.ext
  e.g.  AP_R12_Invoice_Processing_Guide_2024-03-01.pdf

Usage:
  python organizer.py <root_folder> [--dry-run]

Dependencies:
  pip install pdfplumber python-docx openpyxl striprtf
"""

import argparse
import hashlib
import logging
import re
import shutil
import sys
from datetime import datetime
from pathlib import Path

# ── optional dependencies (graceful degradation) ──────────────────────────────

try:
    import pdfplumber
    HAS_PDF = True
except ImportError:
    HAS_PDF = False

try:
    from docx import Document as DocxDocument
    HAS_DOCX = True
except ImportError:
    HAS_DOCX = False

try:
    import openpyxl
    HAS_XLSX = True
except ImportError:
    HAS_XLSX = False

try:
    from striprtf.striprtf import rtf_to_text
    HAS_RTF = True
except ImportError:
    HAS_RTF = False



# ── folder / classification map ───────────────────────────────────────────────

STRUCTURE = {
    "EPM":             ["FCCS", "ARCS", "PBCS", "Free_Form", "Other"],
    "R12_ERP":         ["P2P", "O2C", "R2R", "Security",
                        "CST", "PA", "FA", "eTax", "HMRC", "BIP_Reports", "Other"],
    "Fusion_Cloud_ERP":["P2P", "O2C", "R2R", "Security",
                        "CST", "PA", "FA", "eTax", "HMRC", "BIP_Reports", "Other"],
    "FDI":             ["ERP", "HCM", "SCM", "CX", "EPM", "Other"],
    "Other":           ["Data_Models", "Unclassified"],
}

# Keywords that map to (top_folder, sub_folder)
# Checked against filename + first ~2000 chars of content (case-insensitive)
MODULE_KEYWORDS = [
    # FDI — Oracle Fusion Data Intelligence (must come before ERP/EPM rules)
    (["fusion data intelligence", r"\bfdi\b", "data intelligence platform",
      "fdr tables", "semantic model lineage", "metric calculation logic",
      "data augmentation scripts", "fusion analytics",
      "fusion.*analytics", "analytics.*business.*questions",
      "analytics.*diagrams", "analytics.*subject.*areas",
      "analytics.*tables"], "FDI", None),

    # EPM products
    (["fccs", "financial consolidation", "close cloud"], "EPM", "FCCS"),
    (["arcs", "account reconciliation"], "EPM", "ARCS"),
    (["pbcs", "planning and budgeting", "planning cloud"], "EPM", "PBCS"),
    (["free form", "freeform", "free-form"], "EPM", "Free_Form"),

    # ERP modules — version detected separately
    # HMRC — HR, Payroll, Workforce
    (["payroll", "workforce", "human resources", r"\bhr\b", r"\bhcm\b",
      "benefits guide", "labor distribution", "employee", "workforce deployment",
      "workforce development", "payroll interface", "compensation",
      "absence management", "talent management", "time and labor"], None, "HMRC"),

    # eTax — tax docs not clearly tied to P2P or O2C
    # Only catches docs where tax is the primary topic with no AP/PO/AR/Order context
    (["tax regime", "tax rule", "tax configuration", "tax setup", "tax classification",
      "fiscal classification", "tax exemption", "withholding tax setup",
      "tax determination", r"\betax\b", "e-tax", "transaction tax",
      "tax overview", "tax framework", "tax reporting", r"\btrx tax\b",
      "tax rate", "tax code setup", "tax authority", "configuring tax"], None, "eTax"),

    # P2P — Procure to Pay: AP, PO, Procurement, AP-related Cash Management
    (["accounts payable", "payables", "invoice approval", "invoice processing",
      "supplier invoice", "supplier payment", "supplier site", "supplier bank",
      "searching supplier", "payment process", "payment file", "modify payment",
      "withholding tax", "isupplier", "spreadsheet invoice", "ap inv",
      "payables dashboard", "payables implementation", "payables essential",
      "payables payment", "invoice recognition", "ofc ap", "electronic invoic",
      "purchase order", r"\bpo\b", "iprocurement", "procurement", "requisition",
      "procure to pay", "p2p", "ofc p2p",
      "receiving", "approval setup", "document approval",
      "business unit.*scm", "scm.*procurement",
      "cash management", "bank reconciliation", "bank statement", "bank account",
      "external bank", "payments overview", "cash mgt", "automatic reconciliation",
      "roles.*po", "setups.*po", "enterprise structure.*erp", "erp.*enterprise structure",
      "self service receipt", "selfservice receipt", "self-service receipt",
      r"\bsuppliers?\b"], None, "P2P"),
    # O2C — Order to Cash: AR, Order Entry, Order Management, Cash receipts
    (["accounts receivable", "order to cash", r"\bo2c\b", "order entry",
      "order management", "customer receipt", "cash receipt",
      "customer payment", "credit management", "collections",
      "customer invoice", "customer receipt", "customer account",
      "customer master", "customer site", "autoinvoice", "auto invoice",
      "revenue recognition", "billing", r"\bar\b"], None, "O2C"),
    # R2R — Record to Report: GL, COA, CVR, Enterprise Structure, AHCS, FAH, SLA
    (["general ledger", r"\bgl\b", "chart of accounts", r"\bcoa\b",
      "coa hierarchy", "coa hierarchies", "gl hierarchy", "gl hierarchies",
      "gl smartview", "smartview gl", "smartview connection",
      "journal entry", "journal import", "journal line", "journal approval",
      "journal reversal", "journal balance",
      "cross validation", r"\bcvr\b", "segment value", "segment security",
      "value set", "flexfield", "account combination",
      "enterprise structure", "ledger setup", "ledger config",
      "accounting hub", "ahcs", "fah", "financial accounting hub",
      "subledger accounting", r"\bsla\b", "subledger", "record to report",
      r"\br2r\b", "period close", "intercompany", "consolidat",
      "gl period", "accounting period", "trial balance",
      "gl translation", "currency revaluation", "revaluation",
      "values and hierarchies", "segment values and hierarchies",
      "gl values", "gl hierarchies",
      "legal entity", "legal entities", "enterprise structure",
      "business unit", "ledger set", "primary ledger",
      "secondary ledger", "reporting currency"], None, "R2R"),
    (["user role", "responsibility", "access control", "profile option", "system admin",
      "roles and setup", "role setup", r"\brbac\b", "security setup",
      r"\broles\b", "user management", "user access",
      "role assignment", "role definition", "job role", "duty role",
      "data role", "abstract role", "privilege", "user provisioning",
      "manage users", "create user", "assign role"], None, "Security"),
    (["cost management", "costing", "inventory valuation", "cost accounting",
      "cycle count", "physical inventory", "inventory count", "inventory guide",
      "inventory", "stock", "item cost", "material cost"], None, "CST"),
    (["project accounting", "project costing", "project billing", "project planning",
      "project financ", "project perf", "project revenue", "ppm office",
      "projectcost", "projectrev", "projectperf", "optimize.*project",
      r"\bppm\b", "project portfolio", "project management",
      "project execution", "project contract"], None, "PA"),
    (["fixed assets", "asset management", "depreciation", "asset book"], None, "FA"),
    # BIP Reports — BI, OTBI, BIP docs not tied to a specific module
    (["business intelligence", r"\bbip\b", r"\botbi\b", "bi report", "bi publisher",
      "oracle analytics", "analytics report", "report design", "data model",
      "bursting", "rtf template", "xsl template", "bi query",
      "oracle transactional business intelligence", "fusion analytics",
      "bi cloud", "bi catalog", "subject area", "obiee"], None, "BIP_Reports"),
]

VERSION_KEYWORDS = {
    "R12":    ["r12", "release 12", "ebs 12", "e-business suite 12"],
    "Fusion": ["fusion", "cloud erp", "oracle cloud", "saas"],
    "R11i":   ["r11i", "11i", "release 11", "11.5", "11.0"],
    "EPM":    ["epm", "hyperion", "planning cloud", "fccs", "arcs", "pbcs"],
}

SQL_EXTENSIONS  = {".sql"}
DOC_EXTENSIONS  = {".pdf", ".docx", ".doc", ".txt", ".rtf", ".xlsx", ".xls",
                   ".csv", ".ppt", ".pptx", ".msg", ".eml", ".zip", ".7z",
                   ".rar", ".gz", ".xml", ".json", ".html", ".htm"}
ALL_EXTENSIONS  = SQL_EXTENSIONS | DOC_EXTENSIONS


# ── text extraction & title detection ────────────────────────────────────────

# Boilerplate patterns to reject as titles
_BOILERPLATE = re.compile(
    r"(table of contents|how to use|powerpoint presentation|audio for|"
    r"zoom|webcast will|copyright|all rights reserved|oracle confidential|"
    r"safe harbor|document display|doc display|http[s]?://|www\.|"
    r"support\.oracle|slide \d+|page \d+|\d+ of \d+|^\d+$|"
    r"oracle corporation|click here|presented by|agenda|introduction|"
    r"note there is no|separate dial|applies to|goal|solution|references|"
    r"symptoms|changes|cause|select yes when|prompted to join|placed on hold|"
    r"live webcast|join the audio|computer speakers|dial.in number|"
    r"questions will be taken|q&a|qanda|via the q|replay and the|"
    r"dashboard knowledge|service requests|patches.*updates|community|"
    r"session using|functionality\.|^january|^february|^march|^april|^may|"
    r"^june|^july|^august|^september|^october|^november|^december|"
    r"^microsoft|^q\d|^\d{4}$|\.pptx$|\.docx$|\.xlsx$|\.pdf$|"
    r"pillar.*fa code|functional area tables)",
    re.IGNORECASE
)


def _is_good_title(text: str) -> bool:
    """Return True if text is a plausible document title."""
    text = text.strip()
    if len(text) < 6 or len(text) > 160:
        return False
    if _BOILERPLATE.search(text):
        return False
    alpha = sum(c.isalpha() for c in text)
    if alpha < len(text) * 0.45:
        return False
    return True


def _pdf_title_by_font(path: Path) -> str:
    """
    Find the title of a PDF by:
    1. Metadata title (strip Microsoft PowerPoint / Word prefixes)
    2. Largest font text in the top 50% of page 1
    3. First good text line on page 1
    """
    try:
        with pdfplumber.open(path) as pdf:
            # 1. metadata title — strip common app prefixes
            meta = pdf.metadata or {}
            meta_title = (meta.get("Title") or "").strip()
            meta_title = re.sub(r"^(Microsoft (PowerPoint|Word|Excel)\s*[-–]?\s*)", "", meta_title, flags=re.IGNORECASE).strip()
            if _is_good_title(meta_title):
                return meta_title

            if not pdf.pages:
                return ""

            page = pdf.pages[0]
            page_height = page.height or 1

            # 2. largest font in TOP 50% of page (where titles live)
            chars = [c for c in (page.chars or []) if float(c.get("top", page_height)) < page_height * 0.55]
            if chars:
                from collections import defaultdict
                by_size = defaultdict(list)
                for ch in chars:
                    size = round(float(ch.get("size", 0)), 1)
                    by_size[size].append(ch)

                for size in sorted(by_size.keys(), reverse=True):
                    # group chars into lines by vertical position (within 3pt)
                    size_chars = sorted(by_size[size], key=lambda c: (round(float(c["top"]) / 3), c["x0"]))
                    from itertools import groupby
                    for _, grp in groupby(size_chars, key=lambda c: round(float(c["top"]) / 3)):
                        line = "".join(c["text"] for c in grp).strip()
                        if _is_good_title(line):
                            return line

            # 3. fallback: scan all text lines on page 1
            page_text = page.extract_text() or ""
            for line in page_text.splitlines():
                line = line.strip()
                if _is_good_title(line):
                    return line

    except Exception:
        pass
    return ""


def _local_title_from_text(text: str, module: str = "") -> str:
    """
    Locally derive a meaningful title from document text without sending data anywhere.

    Strategy:
    1. Look for lines that match known Oracle document title patterns
       (e.g. "How To...", "Setting Up...", "Overview of...", "Configuring...")
    2. Find the most content-rich short line in the top portion of the text
    3. Fall back to extracting key noun phrases
    """
    if not text:
        return ""

    lines = [l.strip() for l in text.splitlines() if l.strip()]

    # Pattern 1: lines that look like document titles (action/topic phrases)
    title_patterns = re.compile(
        r"^(how to|setting up|setup|configuring|overview of|introduction to|"
        r"understanding|managing|implementing|using|working with|guide to|"
        r"troubleshooting|best prac|frequently asked|faq|what is|"
        r"creating|defining|processing|running|integrating|migrating)",
        re.IGNORECASE
    )

    # Check first 20 lines for a title-like phrase
    for line in lines[:20]:
        if title_patterns.match(line) and _is_good_title(line):
            return line

    # Pattern 2: lines in Title Case or ALL CAPS that are descriptive
    for line in lines[:15]:
        if not _is_good_title(line):
            continue
        words = line.split()
        if 3 <= len(words) <= 15:
            # title case: most words start with uppercase
            upper_ratio = sum(1 for w in words if w and w[0].isupper()) / len(words)
            if upper_ratio >= 0.6:
                return line

    # Pattern 3: longest meaningful line in first 10 lines
    candidates = [l for l in lines[:10] if _is_good_title(l) and 4 <= len(l.split()) <= 15]
    if candidates:
        return max(candidates, key=len)

    return ""


def extract_title(path: Path, module: str = "", version: str = "") -> str:
    """
    Extract the actual document title from content/metadata.
    Priority order:
      1. Claude API — reads first page, identifies or generates title
      2. PDF largest-font heuristic / Word heading / Excel sheet name
      3. Cleaned original filename
    """
    ext = path.suffix.lower()
    title = ""

    # Step 1: extract raw first-page text for Claude and heuristics
    first_page = ""
    try:
        if ext == ".pdf" and HAS_PDF:
            with pdfplumber.open(path) as pdf:
                first_page = (pdf.pages[0].extract_text() or "") if pdf.pages else ""
        elif ext in (".docx", ".doc") and HAS_DOCX:
            doc = DocxDocument(path)
            first_page = " ".join(p.text for p in doc.paragraphs[:20])
        elif ext in (".xlsx", ".xls") and HAS_XLSX:
            wb = openpyxl.load_workbook(path, read_only=True, data_only=True)
            ws = wb.worksheets[0]
            rows = []
            for row in ws.iter_rows(max_row=5, values_only=True):
                rows.append(" ".join(str(c) for c in row if c))
            first_page = " ".join(rows)
        elif ext in (".txt", ".sql"):
            first_page = path.read_text(encoding="utf-8", errors="ignore")[:1000]
    except Exception:
        pass

    # Step 2: smart local title analysis from first-page text
    if first_page.strip():
        title = _local_title_from_text(first_page, module)
        if title:
            return title

    try:
        if ext == ".pdf" and HAS_PDF:
            title = _pdf_title_by_font(path)

        elif ext in (".docx", ".doc") and HAS_DOCX:
            doc = DocxDocument(path)
            # try document properties title
            meta_title = (doc.core_properties.title or "").strip()
            if _is_good_title(meta_title):
                title = meta_title
            else:
                # find first heading-style paragraph
                for para in doc.paragraphs[:20]:
                    if para.style and "heading" in para.style.name.lower() and _is_good_title(para.text):
                        title = para.text.strip()
                        break
                # fallback: first good paragraph
                if not title:
                    for para in doc.paragraphs[:10]:
                        if _is_good_title(para.text):
                            title = para.text.strip()
                            break

        elif ext in (".xlsx", ".xls") and HAS_XLSX:
            wb = openpyxl.load_workbook(path, read_only=True, data_only=True)
            ws = wb.worksheets[0]
            sheet_name = (ws.title or "").strip()
            if _is_good_title(sheet_name):
                title = sheet_name
            else:
                for row in ws.iter_rows(max_row=5, values_only=True):
                    for cell in row:
                        val = str(cell).strip() if cell else ""
                        if _is_good_title(val):
                            title = val
                            break
                    if title:
                        break

        elif ext in (".txt", ".sql"):
            lines = path.read_text(encoding="utf-8", errors="ignore").splitlines()
            for line in lines[:20]:
                line = line.strip().lstrip("--#/*").strip()
                if _is_good_title(line):
                    title = line
                    break

    except Exception:
        pass

    # fallback: derive from original filename by stripping known prefixes/dates
    if not title:
        stem = path.stem
        sep = r"[\s_]+"
        all_modules = r"(?:AP_PO|P2P|O2C|R2R|GL|COA|CVR|SLA|AHCS|FAH|FA|Security|CST|PA|eTax|HMRC|BIP_Reports|FDI|FCCS|ARCS|PBCS|Unclassified)"
        all_versions = r"(?:Fusion|R12|EPM|R11i|FDI)"
        for _ in range(3):
            stem = re.sub(r"^(?:OFC" + sep + r")?" + all_modules + sep + all_versions + sep, "", stem, flags=re.IGNORECASE)
            stem = re.sub(r"^(?:OFC" + sep + r")?" + all_modules + sep, "", stem, flags=re.IGNORECASE)
        stem = re.sub(r"[\s_]\d{4}[\s\-_]\d{2}[\s\-_]\d{2}(\s+\d+)?$", "", stem)
        stem = re.sub(r"[\s_\-]+", " ", stem).strip()
        title = stem

    return title


def extract_text(path: Path, max_chars: int = 3000) -> str:
    """Extract a sample of text from the file for classification."""
    ext = path.suffix.lower()

    try:
        if ext == ".pdf" and HAS_PDF:
            with pdfplumber.open(path) as pdf:
                text = ""
                for page in pdf.pages[:3]:
                    text += (page.extract_text() or "")
                    if len(text) >= max_chars:
                        break
            return text[:max_chars]

        if ext in (".docx", ".doc") and HAS_DOCX:
            doc = DocxDocument(path)
            text = " ".join(p.text for p in doc.paragraphs[:50])
            return text[:max_chars]

        if ext in (".xlsx", ".xls") and HAS_XLSX:
            wb = openpyxl.load_workbook(path, read_only=True, data_only=True)
            text = ""
            for ws in wb.worksheets[:2]:
                for row in ws.iter_rows(max_row=20, values_only=True):
                    text += " ".join(str(c) for c in row if c) + " "
                    if len(text) >= max_chars:
                        break
            return text[:max_chars]

        if ext == ".rtf" and HAS_RTF:
            raw = path.read_bytes().decode("utf-8", errors="ignore")
            return rtf_to_text(raw)[:max_chars]

        if ext in (".txt", ".sql", ".csv", ".xml", ".json", ".html", ".htm"):
            return path.read_text(encoding="utf-8", errors="ignore")[:max_chars]

        # zip/archive files — classify by filename only, no content extraction
        if ext in (".zip", ".7z", ".rar", ".gz"):
            return ""

    except Exception as e:
        logging.warning(f"Could not extract text from {path.name}: {e}")

    return ""


# ── classification ────────────────────────────────────────────────────────────

# User guide keywords — docs that are official Oracle guides/manuals
_USER_GUIDE_PATTERNS = re.compile(
    r"(user guide|user's guide|implementation guide|admin guide|"
    r"administering oracle|using oracle|getting started|setup guide|"
    r"configuration guide|installation guide|developer guide|"
    r"reference guide|api guide|security guide|upgrade guide|"
    r"migration guide|release guide|whats new|what's new|"
    r"new features|release notes|patch notes)",
    re.IGNORECASE
)

# Oracle Fusion release codes: 25A, 25B, 25C, 26A, 26B, 26C, 24D etc.
_RELEASE_PATTERN = re.compile(r"\b(2[3-9][A-D])\b", re.IGNORECASE)


def detect_release(text: str) :
    """Detect Oracle Fusion release code (e.g. 25A, 26A) from text."""
    match = _RELEASE_PATTERN.search(text)
    return match.group(1).upper() if match else None


def is_user_guide(haystack: str) -> bool:
    """Return True if the document appears to be an Oracle user/admin guide."""
    return bool(_USER_GUIDE_PATTERNS.search(haystack))


def detect_version(text: str) :
    tl = text.lower()
    for version, keywords in VERSION_KEYWORDS.items():
        for kw in keywords:
            if re.search(kw, tl):
                return version
    return None


def classify(path: Path, content: str) :
    """
    Return (top_folder, sub_folder).
    SQL files always go to Other/Data_Models.
    """
    if path.suffix.lower() in SQL_EXTENSIONS:
        return "Other", "Data_Models"

    # include stem with underscores replaced by spaces so OFC_P2P matches "p2p"
    stem_spaced = path.stem.replace("_", " ")
    haystack = (path.stem + " " + stem_spaced + " " + content).lower()
    version  = detect_version(haystack)

    matched_module = None
    for keywords, top_hint, module in MODULE_KEYWORDS:
        for kw in keywords:
            if re.search(kw, haystack):
                matched_module = (top_hint, module)
                break
        if matched_module:
            break

    if not matched_module:
        return "Other", "Unclassified"

    top_hint, module = matched_module

    # eTax — check if tax doc has P2P or O2C context; if so, send there instead
    if module == "eTax":
        p2p_tax = ["accounts payable", "payables", "supplier", "purchase order",
                   r"\bpo\b", "procurement", "withholding tax", "ap tax"]
        o2c_tax = ["accounts receivable", "customer", "sales order", "order to cash",
                   "billing", "revenue", r"\bar\b"]
        if any(re.search(k, haystack) for k in p2p_tax):
            module = "P2P"
        elif any(re.search(k, haystack) for k in o2c_tax):
            module = "O2C"

    # FDI — detect subfolder from filename/content (check filename first for zip files)
    if top_hint == "FDI":
        fname = path.stem.lower()
        if re.search(r"_hcm_|hcm analytics|\bhcm\b", fname):
            return "FDI", "HCM"
        if re.search(r"_cx_|cx analytics|\bcx\b", fname):
            return "FDI", "CX"
        if re.search(r"_scm_|scm analytics|\bscm\b", fname):
            return "FDI", "SCM"
        if re.search(r"_erp_|erp analytics|\berp\b", fname):
            return "FDI", "ERP"
        if re.search(r"\b(erp|fin|gl|ap|ar|p2p|r2r)\b", haystack):
            return "FDI", "ERP"
        if re.search(r"\b(hcm|hr|workforce|payroll)\b", haystack):
            return "FDI", "HCM"
        if re.search(r"\b(scm|supply chain|inventory|warehouse)\b", haystack):
            return "FDI", "SCM"
        if re.search(r"\b(cx|customer experience|sales|service)\b", haystack):
            return "FDI", "CX"
        if re.search(r"\b(epm|planning|budgeting|fccs|pbcs)\b", haystack):
            return "FDI", "EPM"
        return "FDI", "Other"

    # EPM modules always go under EPM
    if top_hint == "EPM":
        return "EPM", module

    # R2R — split into AHCS, GL, Other subfolders
    if module == "R2R":
        r2r_sub = _r2r_module_tag(haystack)
        sub_path = f"R2R/{r2r_sub}"
        if is_user_guide(haystack):
            release = detect_release(path.stem + " " + content)
            release_folder = release if release else "General"
            sub_path = f"R2R/{r2r_sub}/User Guides/{release_folder}"
        if version == "R12" or version == "R11i":
            return "R12_ERP", sub_path
        return "Fusion_Cloud_ERP", sub_path

    # ERP modules — check for user guides before routing by version
    def _route_with_guide(top: str, mod: str) :
        if is_user_guide(haystack):
            release = detect_release(path.stem + " " + content)
            release_folder = release if release else "General"
            return top, f"{mod}/User Guides/{release_folder}"
        return top, mod

    # ERP modules — route by version
    if version == "R12" or version == "R11i":
        return _route_with_guide("R12_ERP", module)
    if version == "Fusion" or version == "EPM":
        return _route_with_guide("Fusion_Cloud_ERP", module)

    # version unclear — default to Fusion_Cloud_ERP rather than losing the file
    return _route_with_guide("Fusion_Cloud_ERP", module)


# ── naming ────────────────────────────────────────────────────────────────────

MAX_FILENAME = 100  # characters excluding extension

# Words to shorten in titles (case-insensitive, longest phrases first)
ABBREVIATIONS = [
    # Multi-word phrases first (must come before single-word matches)
    (r"\bBI Publisher\b",                   "BIP"),
    (r"\bProcure to Pay\b",                 "P2P"),
    (r"\bOrder to Cash\b",                  "O2C"),
    (r"\bRecord to Report\b",               "R2R"),
    (r"\bAccounting Hub\b",                 "AHCS"),
    (r"\bSubledger Accounting\b",           "SLA"),
    (r"\bAccounts Payable\b",               "AP"),
    (r"\bAccounts Receivable\b",            "AR"),
    (r"\bGeneral Ledger\b",                 "GL"),
    (r"\bChart of Accounts\b",              "COA"),
    (r"\bCross Validation Rules?\b",        "CVR"),
    (r"\bFixed Assets?\b",                  "FA"),
    (r"\bJournal Entr(y|ies)\b",            "JE"),
    (r"\bPurchase Orders?\b",               "PO"),
    (r"\bSales Orders?\b",                  "SO"),
    (r"\bWork Orders?\b",                   "WO"),
    (r"\bWork in Process\b",                "WIP"),
    (r"\bElectronic Funds Transfer\b",      "EFT"),
    (r"\bElectronic Data Interchange\b",    "EDI"),
    (r"\bMaterial Requirements Planning\b", "MRP"),
    (r"\bBill of Materials\b",              "BOM"),
    (r"\bCost of Goods Sold\b",             "COGS"),
    (r"\bYear to Date\b",                   "YTD"),
    (r"\bMonth to Date\b",                  "MTD"),
    (r"\bQuarter to Date\b",                "QTD"),
    (r"\bEnd of Month\b",                   "EOM"),
    (r"\bUser Defined Codes?\b",            "UDC"),
    (r"\bBusiness Units?\b",                "BU"),
    (r"\bCost Centers?\b",                  "CC"),
    (r"\bNet Book Value\b",                 "NBV"),
    (r"\bValue Added Tax\b",                "VAT"),
    (r"\bBest Practices\b",                 "Best Prac"),
    (r"\bWhite Paper\b",                    "WP"),
    (r"\bShop Floor Control\b",              "SFC"),
    (r"\bWarehouse Management\b",           "WM"),
    (r"\bQuality Management\b",             "QM"),
    (r"\bHuman Resources?\b",               "HR"),
    (r"\bThird Party Logistics\b",          "3PL"),
    (r"\bAdvanced Planning and Scheduling\b", "APS"),
    (r"\bAssemble to Order\b",              "ATO"),
    (r"\bAvailable to Promise\b",           "ATP"),
    (r"\bBusiness to Business\b",           "B2B"),
    (r"\bBusiness to Consumer\b",           "B2C"),
    (r"\bBusiness Intelligence\b",          "BI"),
    (r"\bBill of Lading\b",                 "BOL"),
    (r"\bConfigure Price Quote\b",          "CPQ"),
    (r"\bCustomer Relationship Management\b", "CRM"),
    (r"\bCapacity Requirements Planning\b", "CRP"),
    (r"\bConfigure to Order\b",             "CTO"),
    (r"\bEngineer to Order\b",              "ETO"),
    (r"\bFirst In[,]? First Out\b",         "FIFO"),
    (r"\bFree on Board\b",                  "FOB"),
    (r"\bJust[- ]in[- ]Time\b",             "JIT"),
    (r"\bKey Performance Indicators?\b",    "KPI"),
    (r"\bLast In[,]? First Out\b",          "LIFO"),
    (r"\bMaster Production Schedule\b",     "MPS"),
    (r"\bMake to Order\b",                  "MTO"),
    (r"\bMake to Stock\b",                  "MTS"),
    (r"\bReturn Material Authorization\b",  "RMA"),
    (r"\bSupply Chain Management\b",        "SCM"),
    (r"\bUser Defined Fields?\b",           "UDF"),
    (r"\bUnit of Measure\b",                "UOM"),
    (r"\bEnterprise Resource Planning\b",   "ERP"),
    # Single words
    (r"\bPayables\b",                       "AP"),
    (r"\bReceivables\b",                    "AR"),
    (r"\bSubledger\b",                      "SBL"),
    (r"\bProcurement\b",                    "Proc"),
    (r"\bConfiguration\b",                  "Config"),
    (r"\bConfigure\b",                      "Config"),
    (r"\bNotifications?\b",                 "Notif"),
    (r"\bImplementation\b",                 "Impl"),
    (r"\bManagement\b",                     "Mgmt"),
    (r"\bInformation\b",                    "Info"),
    (r"\bEnterprise\b",                     "Ent"),
    (r"\bStructures?\b",                    "Struct"),
    (r"\bPerformance\b",                    "Perf"),
    (r"\bFinancial\b",                      "Fin"),
    (r"\bAccounting\b",                     "Acctg"),
    (r"\bTroubleshooting\b",               "Troubleshoot"),
    (r"\bPublisher\b",                      "Pub"),
    (r"\bTemplates?\b",                     "Tmpl"),
    (r"\bSolutions?\b",                     "Soln"),
    (r"\bOrganization\b",                   "Org"),
    (r"\bAuthorization\b",                  "Auth"),
    (r"\bAdministration\b",                 "Admin"),
    (r"\bApplication\b",                    "App"),
    (r"\bTransactions?\b",                  "Txn"),
    (r"\bReconciliation\b",                 "Recon"),
    (r"\bConsolidation\b",                  "Consol"),
    (r"\bDocuments?\b",                     "Doc"),
    (r"\bIntegration\b",                    "Integ"),
    (r"\bCustomization\b",                  "Custom"),
    (r"\bInventory\b",                      "Inv"),
    (r"\bInvoices?\b",                      "Inv"),
    (r"\bDistribution\b",                   "Dist"),
    (r"\bRequirements?\b",                  "Req"),
    (r"\bSpecifications?\b",                "Spec"),
    (r"\bDepartment\b",                     "Dept"),
    (r"\bReference\b",                      "Ref"),
    (r"\bSequence\b",                       "Seq"),
    (r"\bQuantity\b",                       "Qty"),
    (r"\bAmount\b",                         "Amt"),
    (r"\bApproval\b",                       "Apprvl"),
    (r"\bProcessing\b",                     "Proc"),
    (r"\bOverview\b",                       "Overview"),
]


def _clean(value: str) -> str:
    value = re.sub(r'[\\/:*?"<>|+\-_]', " ", value)   # remove unsafe + special chars
    value = re.sub(r'[–—•©®™€£¥°]', " ", value)        # remove special unicode chars
    value = re.sub(r'[^\x20-\x7E]', " ", value)         # remove any remaining non-ASCII
    value = re.sub(r"\s+", " ", value.strip())           # collapse double spaces
    return value


def _shorten(title: str) -> str:
    """Apply abbreviations to reduce title length."""
    for pattern, replacement in ABBREVIATIONS:
        title = re.sub(pattern, replacement, title, flags=re.IGNORECASE)
    title = re.sub(r"\s+", " ", title.strip())
    return title


def _r2r_module_tag(haystack: str) -> str:
    """For R2R files, detect the specific module tag: AHCS, GL, or Other."""
    h = haystack.lower()
    if any(re.search(k, h) for k in ["accounting hub", "ahcs", r"\bfah\b", "financial accounting hub"]):
        return "AHCS"
    if any(re.search(k, h) for k in ["general ledger", r"\bgl\b", "journal", "chart of accounts",
                                      r"\bcoa\b", "coa hierarch", "gl hierarch", "gl smartview",
                                      "smartview gl", "cross validation", r"\bcvr\b",
                                      "subledger", r"\bsla\b", "ledger setup",
                                      "smartview", "essbase", "segment", "hierarchy",
                                      "value set", "flexfield", "account combination",
                                      "trial balance", "gl translation", "revaluation",
                                      "accounting period", "period close",
                                      "values and hierarchies", "gl values", "gl hierarchies",
                                      "legal entity", "legal entities", "enterprise structure",
                                      "business unit", "ledger set", "primary ledger",
                                      "secondary ledger", "reporting currency"]):
        return "GL"
    return "Other"


def build_filename(path: Path, top: str, sub: str, haystack: str = "") -> str:
    """Build new filename: MODULE VERSION Title YYYY-MM-DD.ext, max 100 chars.
    Title is extracted from document content/metadata — not derived from old filename."""
    ext = path.suffix.lower()

    # derive version tag
    version_tag = {
        "R12_ERP":          "R12",
        "Fusion_Cloud_ERP": "Fusion",
        "EPM":              "EPM",
        "FDI":              "FDI",
        "Other":            "",
    }.get(top, "")

    # R2R subfolders — use just the last part as the module tag (AHCS, GL, Other)
    if "/" in sub:
        module_tag = sub.split("/")[-1]
        if module_tag == "Other":
            module_tag = "R2R"
    else:
        module_tag = sub if sub not in ("Other", "Data_Models", "Unclassified") else ""

    # file creation/modification date
    mtime = datetime.fromtimestamp(path.stat().st_mtime).strftime("%Y-%m-%d")

    # extract title from document content — ignore old filename entirely
    raw_title = extract_title(path, module=module_tag, version=version_tag)
    title = _shorten(_clean(raw_title))

    # build full name and enforce 80-char limit on the title portion
    prefix = " ".join(p for p in [module_tag, version_tag] if p)
    suffix = mtime
    # available space for title = MAX_FILENAME - prefix - suffix - 2 spaces
    reserved = len(prefix) + len(suffix) + (2 if prefix else 0) + 1
    max_title = MAX_FILENAME - reserved
    if len(title) > max_title:
        title = title[:max_title].rsplit(" ", 1)[0]  # trim at word boundary

    # detect release code for user guides and prepend to filename
    release_prefix = ""
    if "User Guides" in sub:
        release = detect_release(path.stem)
        if release:
            release_prefix = release

    parts = [p for p in [release_prefix, module_tag, version_tag, title, mtime] if p]
    return " ".join(parts) + ext


# ── move logic ────────────────────────────────────────────────────────────────

def _file_hash(path: Path) -> str:
    """Return MD5 hash of file contents for duplicate detection."""
    h = hashlib.md5()
    with open(path, "rb") as f:
        for chunk in iter(lambda: f.read(65536), b""):
            h.update(chunk)
    return h.hexdigest()


def _build_hash_index(root: Path, inbox: Path) -> dict:
    """Build a hash → path index of all files already in the destination folders.
    Excludes inbox and Other/Unclassified so those files are always re-evaluated."""
    unclassified = str(root / "Other" / "Unclassified")
    duplicates = str(root / "_Duplicates")
    index = {}
    for path in root.rglob("*"):
        if not path.is_file():
            continue
        path_str = str(path)
        if str(inbox) in path_str:
            continue
        if unclassified in path_str:
            continue
        if duplicates in path_str:
            continue
        if path.suffix.lower() == ".txt":
            continue
        try:
            index[_file_hash(path)] = path
        except Exception:
            pass
    return index


def _delete_empty_dirs(folder: Path):
    """Recursively delete empty directories bottom-up inside folder."""
    for dirpath in sorted(folder.rglob("*"), reverse=True):
        if dirpath.is_dir() and dirpath != folder:
            try:
                dirpath.rmdir()  # only succeeds if empty
            except OSError:
                pass  # not empty, skip


def process_inbox(root: Path, dry_run: bool):
    # support both _Inbox and INBOX folder names
    inbox = root / "INBOX" if (root / "INBOX").exists() else root / "_Inbox"
    if not inbox.exists():
        inbox.mkdir(parents=True)
        print(f"Created _Inbox at {inbox} -- add files there and re-run.")
        return

    # recursively find all supported files in inbox and subfolders
    files = [f for f in inbox.rglob("*") if f.is_file() and f.suffix.lower() in ALL_EXTENSIONS]

    # also scan Other/Unclassified so previously unclassified files get retried
    unclassified_dir = root / "Other" / "Unclassified"
    if unclassified_dir.exists():
        unclassified_files = [f for f in unclassified_dir.rglob("*")
                              if f.is_file() and f.suffix.lower() in ALL_EXTENSIONS]
        if unclassified_files:
            print(f"  Also scanning Other/Unclassified ({len(unclassified_files)} file(s))...")
            files = list(files) + unclassified_files

    if not files:
        print(f"No supported files found in {inbox}")
        return

    log_lines = []
    print(f"{'DRY RUN -- ' if dry_run else ''}Processing {len(files)} file(s) from _Inbox (incl. subfolders)\n")

    # build hash index of files already in destination folders
    print("  Building duplicate index...")
    hash_index = _build_hash_index(root, inbox)
    print(f"  {len(hash_index)} existing files indexed.\n")

    for path in sorted(files):
        # show relative path from root so source folder is visible
        try:
            rel = path.relative_to(inbox)
        except ValueError:
            rel = path.relative_to(root)
        print(f"  {rel}")
        content  = extract_text(path)
        top, sub = classify(path, content)
        stem_spaced = path.stem.replace("_", " ")
        haystack = (path.stem + " " + stem_spaced + " " + content).lower()
        new_name = build_filename(path, top, sub, haystack)
        dest_dir = root / top / sub
        dest     = dest_dir / new_name

        # duplicate check — hash the incoming file
        try:
            incoming_hash = _file_hash(path)
        except PermissionError:
            print(f"    [SKIPPED] File locked (OneDrive syncing or open): {path.name}")
            log_lines.append(f"SKIPPED (locked): {rel}")
            continue

        if incoming_hash in hash_index:
            existing = hash_index[incoming_hash]
            print(f"    [DUPLICATE] Already exists as: {existing.relative_to(root)}")
            log_lines.append(f"DUPLICATE: {rel}  ==  {existing.relative_to(root)}")
            if not dry_run:
                try:
                    dup_dir = root / "_Duplicates"
                    dup_dir.mkdir(exist_ok=True)
                    dup_dest = dup_dir / path.name
                    counter = 1
                    while dup_dest.exists():
                        dup_dest = dup_dir / f"{path.stem} {counter}{path.suffix.lower()}"
                        counter += 1
                    shutil.move(str(path), str(dup_dest))
                except PermissionError:
                    print(f"    [SKIPPED] Duplicate locked (OneDrive syncing): {path.name}")
                    log_lines.append(f"SKIPPED (locked duplicate): {rel}")
            continue

        print(f"    -> {top}/{sub}/{dest.name}")
        log_lines.append(f"{rel}  ->  {top}/{sub}/{dest.name}")

        if not dry_run:
            try:
                dest_dir.mkdir(parents=True, exist_ok=True)
                shutil.move(str(path), str(dest))
                hash_index[incoming_hash] = dest  # update index with newly moved file
            except PermissionError:
                print(f"    [SKIPPED] File locked (OneDrive syncing or open): {path.name}")
                log_lines.append(f"SKIPPED (locked): {rel}")

    # delete empty folders left behind in inbox
    if not dry_run:
        _delete_empty_dirs(inbox)
        print(f"\nEmpty inbox subfolders removed.")

    # write log to _Logs folder
    logs_dir = root / "_Logs"
    if not dry_run:
        logs_dir.mkdir(exist_ok=True)
        log_path = logs_dir / f"organizer_log_{datetime.now().strftime('%Y%m%d_%H%M%S')}.txt"
        summary  = "\n".join(log_lines)
        log_path.write_text(summary, encoding="utf-8")
        print(f"\nLog saved to _Logs/{log_path.name}")

        # delete logs older than 1 month
        cutoff = datetime.now().timestamp() - (30 * 24 * 60 * 60)
        deleted = 0
        for old_log in logs_dir.glob("organizer_log_*.txt"):
            if old_log.stat().st_mtime < cutoff:
                old_log.unlink()
                deleted += 1
        if deleted:
            print(f"Deleted {deleted} log(s) older than 30 days.")
    else:
        print(f"\nDry run complete -- no files moved.")

    print(f"\n{len(files)} file(s) processed.")


# ── setup ─────────────────────────────────────────────────────────────────────

def setup_folders(root: Path):
    """Create the full folder structure if it doesn't exist."""
    for top, subs in STRUCTURE.items():
        for sub in subs:
            (root / top / sub).mkdir(parents=True, exist_ok=True)
    (root / "_Inbox").mkdir(exist_ok=True)
    (root / "_Logs").mkdir(exist_ok=True)
    print(f"Folder structure ready at: {root}")


# ── CLI ───────────────────────────────────────────────────────────────────────

def main():
    parser = argparse.ArgumentParser(description="Oracle document organizer.")
    parser.add_argument("root", help="Root folder where Oracle Docs structure lives")
    parser.add_argument("--dry-run", action="store_true", help="Preview moves without changing anything")
    parser.add_argument("--setup", action="store_true", help="Create folder structure and exit")
    args = parser.parse_args()

    root = Path(args.root).expanduser().resolve()
    root.mkdir(parents=True, exist_ok=True)

    if args.setup:
        setup_folders(root)
        return

    process_inbox(root, dry_run=args.dry_run)


if __name__ == "__main__":
    logging.basicConfig(level=logging.WARNING)
    main()
